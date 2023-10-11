# 14. What integers represent the levels of headings available in Word documents?

print('integers 0,1, 2, 3, and 4 used to represent the levels of headings available in Word documents')


import docx
doc = docx.Document()
doc.add_heading('Header 0', 0)
doc.add_heading('Header 1', 1)
doc.add_heading('Header 2', 2)
doc.add_heading('Header 3', 3)
doc.add_heading('Header 4', 4)
doc.save('headings.docx')

