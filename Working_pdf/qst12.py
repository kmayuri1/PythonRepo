# 12. How do you create a Document object for a new Word document?


print('By Calling the docx.Document() function')

import docx
doc=docx.Document('demo.docx')
print(len(doc.paragraphs))
print(doc.paragraphs[0].text)