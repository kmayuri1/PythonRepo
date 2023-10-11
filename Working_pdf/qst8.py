# 8. What is the difference between a Paragraph object and a Run object?

print('The Document object contains a list of Paragraph objects for the paragraphs in the document. Each of these Paragraph objects contains a list of one or more Run objects and  A new Run object is needed whenever the text style changes'
      )


import docx
doc = docx.Document('demo.docx')
len(doc.paragraphs)
print(doc.paragraphs[0].text)
print(doc.paragraphs[0].runs[0].text)

 