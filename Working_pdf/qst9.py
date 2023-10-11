# 9. How do you obtain a list of Paragraph objects for a Document object that’s stored in a variable named doc?
print('Document object, which has a paragraphs attribute that is a list of Paragraph objects.')

import docx
doc=docx.Document('demo.docx')
print(len(doc.paragraphs))
print(doc.paragraphs[0].text)

